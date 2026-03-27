"""Word document creation and editing tool.

Creates structured DOCX documents from markdown text input. Supports
optional template files, custom styling, and outputs the result as
base64-encoded bytes or writes to a local path / Azure Blob Storage.
"""

from __future__ import annotations

import base64
import io
import logging
import pathlib
import re
from typing import Any, Literal

from pydantic import BaseModel, Field

from app.core.tool_registry import define_tool

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style presets
# ---------------------------------------------------------------------------

_STYLE_PRESETS: dict[str, dict[str, Any]] = {
    "corporate": {
        "title_font": "Calibri",
        "body_font": "Calibri",
        "title_size_pt": 28,
        "heading1_size_pt": 22,
        "heading2_size_pt": 18,
        "heading3_size_pt": 14,
        "body_size_pt": 11,
        "title_color": "1E3A5F",
        "heading_color": "1E3A5F",
        "body_color": "333333",
        "line_spacing": 1.15,
    },
    "modern": {
        "title_font": "Segoe UI",
        "body_font": "Segoe UI",
        "title_size_pt": 32,
        "heading1_size_pt": 24,
        "heading2_size_pt": 18,
        "heading3_size_pt": 14,
        "body_size_pt": 11,
        "title_color": "0078D4",
        "heading_color": "0078D4",
        "body_color": "2D2D2D",
        "line_spacing": 1.2,
    },
    "classic": {
        "title_font": "Times New Roman",
        "body_font": "Times New Roman",
        "title_size_pt": 26,
        "heading1_size_pt": 20,
        "heading2_size_pt": 16,
        "heading3_size_pt": 13,
        "body_size_pt": 12,
        "title_color": "000000",
        "heading_color": "000000",
        "body_color": "000000",
        "line_spacing": 1.5,
    },
    "minimal": {
        "title_font": "Helvetica",
        "body_font": "Helvetica",
        "title_size_pt": 28,
        "heading1_size_pt": 20,
        "heading2_size_pt": 16,
        "heading3_size_pt": 13,
        "body_size_pt": 11,
        "title_color": "333333",
        "heading_color": "555555",
        "body_color": "444444",
        "line_spacing": 1.15,
    },
}

# ---------------------------------------------------------------------------
# Parameters
# ---------------------------------------------------------------------------


class CreateDocxParams(BaseModel):
    """Parameters for the create_docx tool."""

    markdown: str = Field(
        description=(
            "Markdown text to convert into a structured Word document. "
            "Supports headings (# ## ### ####), bold (**text**), "
            "italic (*text*), inline code (`code`), bullet lists (- or *), "
            "numbered lists (1. 2. 3.), blockquotes (>), horizontal rules "
            "(---), tables (| col | col |), and code blocks (``` ... ```)."
        ),
    )
    style_preset: Literal[
        "corporate", "modern", "classic", "minimal"
    ] | None = Field(
        default=None,
        description=(
            "Optional style preset to apply. Ignored when a template file "
            "is provided. Available: corporate, modern, classic, minimal."
        ),
    )
    template_bytes_b64: str | None = Field(
        default=None,
        description=(
            "Base64-encoded bytes of a .docx template file. The template's "
            "styles (fonts, colours, spacing) will be preserved and the "
            "markdown content will be written using those styles."
        ),
    )
    template_local_path: str | None = Field(
        default=None,
        description="Local filesystem path to a .docx template file.",
    )
    template_storage_account: str | None = Field(
        default=None,
        description="Azure Storage account name for a template stored in Blob Storage.",
    )
    template_container: str | None = Field(
        default=None,
        description="Azure Blob container holding the template file.",
    )
    template_blob_path: str | None = Field(
        default=None,
        description="Blob path to the .docx template file.",
    )
    output_local_path: str | None = Field(
        default=None,
        description=(
            "If provided, the generated DOCX is saved to this local path "
            "and the response includes the path instead of base64 content."
        ),
    )
    output_storage_account: str | None = Field(
        default=None,
        description="Azure Storage account to upload the generated DOCX to.",
    )
    output_container: str | None = Field(
        default=None,
        description="Azure Blob container for the output DOCX.",
    )
    output_blob_path: str | None = Field(
        default=None,
        description="Blob path for the output DOCX.",
    )
    filename: str = Field(
        default="document.docx",
        description="Filename for the generated document (used in metadata).",
    )


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

# Regex patterns for markdown elements
_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)")
_RE_BULLET = re.compile(r"^(\s*)[-*]\s+(.*)")
_RE_NUMBERED = re.compile(r"^(\s*)\d+\.\s+(.*)")
_RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)")
_RE_HR = re.compile(r"^-{3,}$|^\*{3,}$|^_{3,}$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$")
_RE_CODE_FENCE = re.compile(r"^```")

# Inline formatting
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
_RE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)(.+?)(?<!_)_(?!_)")
_RE_INLINE_CODE = re.compile(r"`([^`]+)`")
_RE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


class _ParsedBlock:
    """Intermediate representation of a markdown block."""

    __slots__ = ("kind", "text", "level", "rows", "language")

    def __init__(
        self,
        kind: str,
        text: str = "",
        level: int = 0,
        rows: list[list[str]] | None = None,
        language: str = "",
    ) -> None:
        self.kind = kind
        self.text = text
        self.level = level
        self.rows = rows or []
        self.language = language


def _parse_markdown(md: str) -> list[_ParsedBlock]:
    """Parse markdown text into a list of block elements."""
    blocks: list[_ParsedBlock] = []
    lines = md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code fence
        if _RE_CODE_FENCE.match(line):
            lang = line.strip("`").strip()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not _RE_CODE_FENCE.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            blocks.append(
                _ParsedBlock("code_block", "\n".join(code_lines), language=lang)
            )
            i += 1
            continue

        # Horizontal rule
        if _RE_HR.match(line.strip()):
            blocks.append(_ParsedBlock("hr"))
            i += 1
            continue

        # Heading
        m = _RE_HEADING.match(line)
        if m:
            level = len(m.group(1))
            blocks.append(_ParsedBlock("heading", m.group(2).strip(), level=level))
            i += 1
            continue

        # Table: collect consecutive table rows
        if _RE_TABLE_ROW.match(line.strip()):
            table_rows: list[list[str]] = []
            while i < len(lines) and _RE_TABLE_ROW.match(lines[i].strip()):
                row_line = lines[i].strip()
                if _RE_TABLE_SEP.match(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            blocks.append(_ParsedBlock("table", rows=table_rows))
            continue

        # Blockquote
        m = _RE_BLOCKQUOTE.match(line)
        if m:
            quote_lines: list[str] = []
            while i < len(lines):
                qm = _RE_BLOCKQUOTE.match(lines[i])
                if qm:
                    quote_lines.append(qm.group(1))
                    i += 1
                else:
                    break
            blocks.append(_ParsedBlock("blockquote", "\n".join(quote_lines)))
            continue

        # Bullet list item
        m = _RE_BULLET.match(line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            blocks.append(_ParsedBlock("bullet", m.group(2).strip(), level=level))
            i += 1
            continue

        # Numbered list item
        m = _RE_NUMBERED.match(line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            blocks.append(_ParsedBlock("numbered", m.group(2).strip(), level=level))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — collect contiguous non-blank, non-special lines
        para_lines: list[str] = []
        while i < len(lines):
            l = lines[i]
            if (
                not l.strip()
                or _RE_HEADING.match(l)
                or _RE_BULLET.match(l)
                or _RE_NUMBERED.match(l)
                or _RE_BLOCKQUOTE.match(l)
                or _RE_HR.match(l.strip())
                or _RE_TABLE_ROW.match(l.strip())
                or _RE_CODE_FENCE.match(l)
            ):
                break
            para_lines.append(l)
            i += 1
        blocks.append(_ParsedBlock("paragraph", " ".join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------


def _hex_to_rgb(hex_str: str) -> tuple[int, int, int]:
    """Convert '1E3A5F' to (30, 58, 95)."""
    h = hex_str.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _apply_inline_formatting(paragraph, text: str) -> None:
    """Parse inline markdown (bold, italic, code, links) and add runs."""
    from docx.shared import RGBColor, Pt

    # Tokenize the text into segments with formatting info
    segments: list[tuple[str, set[str]]] = []
    pos = 0

    # Combined pattern for all inline elements
    combined = re.compile(
        r"(\*\*(.+?)\*\*)"       # bold
        r"|(__(.+?)__)"           # bold alt
        r"|(`([^`]+)`)"          # inline code
        r"|(\[([^\]]+)\]\(([^)]+)\))"  # link
        r"|((?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*))"  # italic
        r"|((?<!_)_(?!_)(.+?)(?<!_)_(?!_))"        # italic alt
    )

    for m in combined.finditer(text):
        # Add any plain text before this match
        if m.start() > pos:
            segments.append((text[pos:m.start()], set()))

        if m.group(2):  # **bold**
            segments.append((m.group(2), {"bold"}))
        elif m.group(4):  # __bold__
            segments.append((m.group(4), {"bold"}))
        elif m.group(6):  # `code`
            segments.append((m.group(6), {"code"}))
        elif m.group(8):  # [text](url)
            segments.append((m.group(8), {"link"}))
        elif m.group(11):  # *italic*
            segments.append((m.group(11), {"italic"}))
        elif m.group(13):  # _italic_
            segments.append((m.group(13), {"italic"}))

        pos = m.end()

    # Trailing plain text
    if pos < len(text):
        segments.append((text[pos:], set()))

    # If no formatting was found, just add the whole text
    if not segments:
        segments = [(text, set())]

    for seg_text, formats in segments:
        run = paragraph.add_run(seg_text)
        if "bold" in formats:
            run.bold = True
        if "italic" in formats:
            run.italic = True
        if "code" in formats:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _build_document(
    blocks: list[_ParsedBlock],
    template_bytes: bytes | None,
    style: dict[str, Any],
) -> bytes:
    """Build a python-docx Document from parsed blocks and return bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    if template_bytes:
        doc = Document(io.BytesIO(template_bytes))
    else:
        doc = Document()

    # Track whether this is a fresh doc (no template) for style customisation
    apply_custom_styles = template_bytes is None

    # ---- configure default styles when no template is used ----
    if apply_custom_styles:
        _configure_styles(doc, style)

    # ---- write blocks ----
    for block in blocks:
        if block.kind == "heading":
            level = min(block.level, 4)
            if level == 1 and apply_custom_styles:
                p = doc.add_heading(block.text, level=0)
                _style_heading(p, style, "title")
            else:
                p = doc.add_heading(block.text, level=level)
                if apply_custom_styles:
                    _style_heading(p, style, f"heading{level}")

        elif block.kind == "paragraph":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, block.text)
            if apply_custom_styles:
                _style_paragraph(p, style)

        elif block.kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, block.text)
            if apply_custom_styles:
                _style_paragraph(p, style)
            if block.level > 0:
                _set_list_indent(p, block.level)

        elif block.kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, block.text)
            if apply_custom_styles:
                _style_paragraph(p, style)
            if block.level > 0:
                _set_list_indent(p, block.level)

        elif block.kind == "blockquote":
            p = doc.add_paragraph()
            _apply_inline_formatting(p, block.text)
            if apply_custom_styles:
                _style_paragraph(p, style)
            # Indent and colour for blockquotes
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        elif block.kind == "code_block":
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            # Light grey background shading
            shading = p.paragraph_format._element.get_or_add_pPr()
            shd = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F5F5F5",
            })
            shading.append(shd)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif block.kind == "table":
            _add_table(doc, block.rows, style, apply_custom_styles)

        elif block.kind == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Thin bottom border as a horizontal rule
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

    # ---- serialise ----
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _configure_styles(doc, style: dict[str, Any]) -> None:
    """Configure document-level default styles."""
    from docx.shared import Pt, RGBColor

    st = doc.styles["Normal"]
    st.font.name = style.get("body_font", "Calibri")
    st.font.size = Pt(style.get("body_size_pt", 11))
    r, g, b = _hex_to_rgb(style.get("body_color", "333333"))
    st.font.color.rgb = RGBColor(r, g, b)
    st.paragraph_format.line_spacing = style.get("line_spacing", 1.15)


def _style_heading(paragraph, style: dict[str, Any], kind: str) -> None:
    """Apply custom font/colour to a heading paragraph."""
    from docx.shared import Pt, RGBColor

    font_name = style.get("title_font", "Calibri")
    if kind == "title":
        size = style.get("title_size_pt", 28)
    elif kind == "heading1":
        size = style.get("heading1_size_pt", 22)
    elif kind == "heading2":
        size = style.get("heading2_size_pt", 18)
    else:
        size = style.get("heading3_size_pt", 14)

    r, g, b = _hex_to_rgb(style.get("heading_color", "1E3A5F"))
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(r, g, b)


def _style_paragraph(paragraph, style: dict[str, Any]) -> None:
    """Apply body font styling to a paragraph's existing runs."""
    from docx.shared import Pt, RGBColor

    r, g, b = _hex_to_rgb(style.get("body_color", "333333"))
    for run in paragraph.runs:
        if run.font.name is None:
            run.font.name = style.get("body_font", "Calibri")
        if run.font.size is None:
            run.font.size = Pt(style.get("body_size_pt", 11))
        if run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(r, g, b)


def _set_list_indent(paragraph, level: int) -> None:
    """Increase left indent for nested list items."""
    from docx.shared import Inches

    paragraph.paragraph_format.left_indent = Inches(0.5 * (level + 1))


def _add_table(
    doc, rows: list[list[str]], style: dict[str, Any], apply_styles: bool
) -> None:
    """Add a formatted table to the document."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if apply_styles:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = style.get("body_font", "Calibri")
                            run.font.size = Pt(style.get("body_size_pt", 11) - 1)

    # Bold header row
    if len(rows) > 0:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        # Header row shading
        heading_color = style.get("heading_color", "1E3A5F")
        for cell in table.rows[0].cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            shd = tcPr.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): heading_color,
            })
            tcPr.append(shd)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Template loader
# ---------------------------------------------------------------------------


async def _load_template(params: CreateDocxParams) -> bytes | None:
    """Resolve template bytes from the various input sources."""
    if params.template_bytes_b64:
        logger.info("Loading template from base64 bytes")
        return base64.b64decode(params.template_bytes_b64)

    if params.template_local_path:
        resolved = pathlib.Path(params.template_local_path).expanduser().resolve()
        if not resolved.is_file():
            raise FileNotFoundError(f"Template file not found: {resolved}")
        logger.info("Loading template from local path %s", resolved)
        return resolved.read_bytes()

    if (
        params.template_storage_account
        and params.template_container
        and params.template_blob_path
    ):
        from app.connectors.blob_connector import BlobConnectorAdapter

        connector = BlobConnectorAdapter(
            account_name=params.template_storage_account,
            credential_type="default_azure_credential",
        )
        try:
            logger.info(
                "Downloading template blob %s/%s",
                params.template_container,
                params.template_blob_path,
            )
            return await connector.download_blob(
                params.template_container, params.template_blob_path
            )
        finally:
            await connector.close()

    return None


# ---------------------------------------------------------------------------
# Output writer
# ---------------------------------------------------------------------------


async def _write_output(
    params: CreateDocxParams, docx_bytes: bytes
) -> dict[str, Any]:
    """Write the generated document to the requested destination."""
    # Local file
    if params.output_local_path:
        out_path = pathlib.Path(params.output_local_path).expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_bytes(docx_bytes)
        logger.info("Saved DOCX to %s (%d bytes)", out_path, len(docx_bytes))
        return {
            "status": "saved",
            "local_path": str(out_path),
            "size_bytes": len(docx_bytes),
            "filename": params.filename,
        }

    # Azure Blob Storage
    if (
        params.output_storage_account
        and params.output_container
        and params.output_blob_path
    ):
        from app.connectors.blob_connector import BlobConnectorAdapter

        connector = BlobConnectorAdapter(
            account_name=params.output_storage_account,
            credential_type="default_azure_credential",
        )
        try:
            await connector.upload_blob(
                container=params.output_container,
                path=params.output_blob_path,
                data=docx_bytes,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            logger.info(
                "Uploaded DOCX to blob %s/%s (%d bytes)",
                params.output_container,
                params.output_blob_path,
                len(docx_bytes),
            )
            return {
                "status": "uploaded",
                "blob_path": f"{params.output_container}/{params.output_blob_path}",
                "size_bytes": len(docx_bytes),
                "filename": params.filename,
            }
        finally:
            await connector.close()

    # Default: return base64
    b64 = base64.b64encode(docx_bytes).decode("ascii")
    return {
        "status": "created",
        "content_b64": b64,
        "size_bytes": len(docx_bytes),
        "filename": params.filename,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


# ---------------------------------------------------------------------------
# Tool definition
# ---------------------------------------------------------------------------


@define_tool(
    name="create_docx",
    description=(
        "Create or edit a Word document (DOCX) from markdown text. Converts "
        "markdown headings, bold, italic, code, bullet/numbered lists, "
        "blockquotes, tables, horizontal rules, and code blocks into "
        "properly formatted Word elements. "
        "Optionally accepts a .docx template file whose styles (fonts, "
        "colours, spacing) are preserved — the markdown content is written "
        "using the template's existing styles. "
        "Without a template, use style_preset (corporate, modern, classic, "
        "minimal) to control appearance. "
        "Output can be returned as base64, saved to a local path, or "
        "uploaded to Azure Blob Storage."
    ),
    parameters_model=CreateDocxParams,
)
async def create_docx(
    params: CreateDocxParams, context: dict
) -> dict[str, Any]:
    """Create a DOCX document from markdown input."""
    try:
        # Resolve template
        template_bytes: bytes | None = None
        try:
            template_bytes = await _load_template(params)
        except FileNotFoundError as e:
            return {"error": str(e), "filename": params.filename}
        except Exception as e:
            return {
                "error": f"Failed to load template: {e}",
                "filename": params.filename,
            }

        # Resolve style preset
        style = _STYLE_PRESETS.get(
            params.style_preset or "corporate", _STYLE_PRESETS["corporate"]
        )

        # Parse and build
        blocks = _parse_markdown(params.markdown)
        logger.info(
            "create_docx: parsed %d blocks, template=%s, style=%s",
            len(blocks),
            "yes" if template_bytes else "no",
            params.style_preset or "corporate",
        )

        docx_bytes = _build_document(blocks, template_bytes, style)

        # Write output
        return await _write_output(params, docx_bytes)

    except Exception as exc:
        logger.exception("create_docx failed for %s", params.filename)
        return {"error": str(exc), "filename": params.filename}
