"""Tool for extracting text, tables, styles, and metadata from DOCX files."""

from __future__ import annotations

import base64
import io
import logging
import pathlib
from typing import Any, Literal

from pydantic import BaseModel, Field

from app.core.tool_registry import define_tool

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Parameters
# ---------------------------------------------------------------------------

class ExtractDocxParams(BaseModel):
    """Parameters for the extract_docx tool."""

    file_bytes: bytes | None = Field(
        None,
        description="Raw bytes of the DOCX file. Preferred when the caller already has the content in memory.",
    )
    file_bytes_b64: str | None = Field(
        None,
        description="Base64-encoded DOCX file content.",
    )
    local_path: str | None = Field(
        None,
        description="Absolute or relative path to a DOCX file on the local filesystem.",
    )
    storage_account: str | None = Field(
        default=None,
        description=(
            "Azure Storage account name to connect to. "
            "Provide this to read from a storage account "
            "(authentication uses DefaultAzureCredential)."
        ),
    )
    container: str | None = Field(
        None, description="Azure Blob Storage container name."
    )
    path: str | None = Field(
        None, description="Blob path to the DOCX file."
    )
    format: Literal["markdown", "json"] = Field(
        "markdown",
        description="Output format: 'markdown' returns a rendered markdown document, 'json' returns structured data.",
    )
    filename: str | None = Field(
        None,
        description="Optional descriptive filename for the result metadata.",
    )


# ---------------------------------------------------------------------------
# Tool definition
# ---------------------------------------------------------------------------

@define_tool(
    name="extract_docx",
    description=(
        "Extract text, tables, styles, and metadata from a Word DOCX file. "
        "Accepts the file as raw bytes (file_bytes), base64-encoded bytes "
        "(file_bytes_b64), a local filesystem path (local_path), or reads "
        "from Azure Blob Storage (storage_account + container + path). Returns a markdown "
        "document (default) or structured JSON. Uses python-docx for parsing."
    ),
    parameters_model=ExtractDocxParams,
)
async def extract_docx(params: ExtractDocxParams, context: dict) -> dict[str, Any]:
    """Extract contents from a DOCX provided as bytes, local path, or fetched from blob storage."""
    try:
        from docx import Document

        # Resolve the file bytes from the supplied source.
        if params.file_bytes:
            logger.info("Parsing DOCX from supplied raw bytes")
            docx_bytes: bytes = params.file_bytes
        elif params.file_bytes_b64:
            logger.info("Parsing DOCX from base64-encoded bytes")
            docx_bytes = base64.b64decode(params.file_bytes_b64)
        elif params.local_path:
            resolved = pathlib.Path(params.local_path).expanduser().resolve()
            if not resolved.is_file():
                return {
                    "error": f"Local file not found: {resolved}",
                    "filename": params.filename or params.local_path,
                }
            logger.info("Reading DOCX from local path %s", resolved)
            docx_bytes = resolved.read_bytes()
        elif params.storage_account and params.container and params.path:
            from app.connectors.blob_connector import BlobConnectorAdapter
            try:
                connector = BlobConnectorAdapter(
                    account_name=params.storage_account,
                    credential_type="default_azure_credential",
                )
                logger.info("Downloading DOCX blob %s/%s", params.container, params.path)
                docx_bytes = await connector.download_blob(params.container, params.path)
            except Exception as blob_err:
                return {
                    "error": f"Failed to download DOCX from blob storage: {blob_err}",
                    "filename": params.filename or params.path,
                }
            finally:
                if connector is not None:
                    await connector.close()
        else:
            return {
                "error": (
                    "Provide one of: file_bytes, file_bytes_b64, local_path, "
                    "or storage_account + container + path."
                ),
                "filename": params.filename or "",
            }

        doc = Document(io.BytesIO(docx_bytes))

        # --- text ---
        text = "\n".join(para.text for para in doc.paragraphs)

        # --- tables ---
        tables: list[dict[str, Any]] = []
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text for cell in row.cells])
            if rows_data:
                headers = rows_data[0]
                data_rows = rows_data[1:]
            else:
                headers = []
                data_rows = []
            tables.append({"headers": headers, "rows": data_rows})

        # --- styles ---
        styles: list[str] = list(
            {para.style.name for para in doc.paragraphs if para.style and para.style.name}
        )

        # --- metadata ---
        metadata: dict[str, Any] = {}
        try:
            props = doc.core_properties
            metadata = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "revision": props.revision,
            }
        except Exception as meta_err:
            logger.warning("Could not extract DOCX metadata: %s", meta_err)

        label = params.filename or params.local_path or params.path or "unknown.docx"

        if params.format == "markdown":
            return {
                "filename": label,
                "content": _docx_to_markdown(label, text, tables, styles, metadata),
            }

        return {
            "filename": label,
            "text": text,
            "tables": tables,
            "styles": styles,
            "metadata": metadata,
        }

    except Exception as exc:
        label = params.filename or params.local_path or params.path or "unknown.docx"
        logger.exception("extract_docx failed for %s", label)
        return {"error": str(exc), "filename": label}


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------

def _docx_to_markdown(
    label: str,
    text: str,
    tables: list[dict[str, Any]],
    styles: list[str],
    metadata: dict[str, Any],
) -> str:
    """Convert extracted DOCX data into a readable markdown document."""
    parts: list[str] = [f"# {label}", ""]

    # Metadata
    meta_items = {k: v for k, v in metadata.items() if v}
    if meta_items:
        parts.append("## Metadata")
        parts.append("")
        for key, value in meta_items.items():
            parts.append(f"- **{key}:** {value}")
        parts.append("")

    # Body text
    if text.strip():
        parts.append("## Content")
        parts.append("")
        parts.append(text.strip())
        parts.append("")

    # Tables
    if tables:
        parts.append("## Tables")
        parts.append("")
        for i, tbl in enumerate(tables, 1):
            parts.append(f"### Table {i}")
            parts.append("")
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers:
                parts.append("| " + " | ".join(str(h) for h in headers) + " |")
                parts.append("| " + " | ".join("---" for _ in headers) + " |")
                for row in rows:
                    parts.append("| " + " | ".join(str(c) for c in row) + " |")
            parts.append("")

    return "\n".join(parts).rstrip() + "\n"
