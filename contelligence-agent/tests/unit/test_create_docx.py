"""Unit tests for the create_docx tool."""

from __future__ import annotations

import base64
import io
import os
import tempfile

import pytest

from app.tools.creation.create_docx import (
    create_docx,
    _parse_markdown,
    _build_document,
    _STYLE_PRESETS,
    CreateDocxParams,
)


# ---------------------------------------------------------------------------
# Markdown parser tests
# ---------------------------------------------------------------------------


class TestParseMarkdown:
    def test_heading_levels(self):
        md = "# Title\n## Section\n### Sub\n#### Deep"
        blocks = _parse_markdown(md)
        headings = [b for b in blocks if b.kind == "heading"]
        assert len(headings) == 4
        assert headings[0].level == 1
        assert headings[0].text == "Title"
        assert headings[1].level == 2
        assert headings[2].level == 3
        assert headings[3].level == 4

    def test_bullet_list(self):
        md = "- Item one\n- Item two\n  - Nested"
        blocks = _parse_markdown(md)
        bullets = [b for b in blocks if b.kind == "bullet"]
        assert len(bullets) == 3
        assert bullets[0].text == "Item one"
        assert bullets[2].level == 1

    def test_numbered_list(self):
        md = "1. First\n2. Second\n3. Third"
        blocks = _parse_markdown(md)
        numbered = [b for b in blocks if b.kind == "numbered"]
        assert len(numbered) == 3

    def test_code_block(self):
        md = "```python\nprint('hello')\n```"
        blocks = _parse_markdown(md)
        code = [b for b in blocks if b.kind == "code_block"]
        assert len(code) == 1
        assert code[0].text == "print('hello')"
        assert code[0].language == "python"

    def test_blockquote(self):
        md = "> This is a quote\n> continued"
        blocks = _parse_markdown(md)
        quotes = [b for b in blocks if b.kind == "blockquote"]
        assert len(quotes) == 1
        assert "This is a quote" in quotes[0].text

    def test_table(self):
        md = "| Name | Age |\n|---|---|\n| Alice | 30 |\n| Bob | 25 |"
        blocks = _parse_markdown(md)
        tables = [b for b in blocks if b.kind == "table"]
        assert len(tables) == 1
        assert len(tables[0].rows) == 3  # header + 2 data rows

    def test_horizontal_rule(self):
        md = "Some text\n\n---\n\nMore text"
        blocks = _parse_markdown(md)
        hrs = [b for b in blocks if b.kind == "hr"]
        assert len(hrs) == 1

    def test_paragraph(self):
        md = "This is a simple paragraph."
        blocks = _parse_markdown(md)
        paras = [b for b in blocks if b.kind == "paragraph"]
        assert len(paras) == 1
        assert paras[0].text == "This is a simple paragraph."


# ---------------------------------------------------------------------------
# Document builder tests
# ---------------------------------------------------------------------------


class TestBuildDocument:
    def test_creates_valid_docx_bytes(self):
        blocks = _parse_markdown("# Hello\n\nWorld")
        result = _build_document(blocks, None, _STYLE_PRESETS["corporate"])
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_docx_is_loadable(self):
        from docx import Document

        blocks = _parse_markdown("# Test\n\n- Item 1\n- Item 2\n\nParagraph here.")
        result = _build_document(blocks, None, _STYLE_PRESETS["modern"])
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_all_style_presets(self):
        blocks = _parse_markdown("# Title\n\nBody text")
        for name, style in _STYLE_PRESETS.items():
            result = _build_document(blocks, None, style)
            assert len(result) > 0, f"Preset {name} produced empty output"

    def test_template_is_used(self):
        from docx import Document

        # Create a minimal template
        template_doc = Document()
        template_doc.add_paragraph("TEMPLATE MARKER")
        buf = io.BytesIO()
        template_doc.save(buf)
        template_bytes = buf.getvalue()

        blocks = _parse_markdown("# New Content\n\nAdded paragraph.")
        result = _build_document(blocks, template_bytes, _STYLE_PRESETS["corporate"])
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "TEMPLATE MARKER" in texts
        assert "New Content" in texts

    def test_table_in_document(self):
        from docx import Document

        md = "| Col A | Col B |\n|---|---|\n| 1 | 2 |"
        blocks = _parse_markdown(md)
        result = _build_document(blocks, None, _STYLE_PRESETS["corporate"])
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "Col A"


# ---------------------------------------------------------------------------
# Full tool integration tests
# ---------------------------------------------------------------------------


class TestCreateDocxTool:
    @pytest.mark.asyncio
    async def test_returns_base64(self):
        params = CreateDocxParams(
            markdown="# Hello World\n\nThis is a test document.",
            filename="test.docx",
        )
        result = await create_docx.handler(params, {})
        assert result["status"] == "created"
        assert "content_b64" in result
        raw = base64.b64decode(result["content_b64"])
        assert len(raw) > 0

    @pytest.mark.asyncio
    async def test_save_to_local_path(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = os.path.join(tmpdir, "output.docx")
            params = CreateDocxParams(
                markdown="# Local Save\n\n- Works great",
                output_local_path=out_path,
                filename="output.docx",
            )
            result = await create_docx.handler(params, {})
            assert result["status"] == "saved"
            assert os.path.isfile(out_path)
            assert os.path.getsize(out_path) > 0

    @pytest.mark.asyncio
    async def test_style_preset(self):
        params = CreateDocxParams(
            markdown="# Styled\n\nBody text with **bold** and *italic*.",
            style_preset="modern",
            filename="styled.docx",
        )
        result = await create_docx.handler(params, {})
        assert result["status"] == "created"

    @pytest.mark.asyncio
    async def test_template_from_b64(self):
        from docx import Document

        # Create template
        template_doc = Document()
        template_doc.add_paragraph("Template content")
        buf = io.BytesIO()
        template_doc.save(buf)
        b64 = base64.b64encode(buf.getvalue()).decode()

        params = CreateDocxParams(
            markdown="# From Template\n\nNew content here.",
            template_bytes_b64=b64,
            filename="from_template.docx",
        )
        result = await create_docx.handler(params, {})
        assert result["status"] == "created"

    @pytest.mark.asyncio
    async def test_complex_markdown(self):
        md = """# Project Report

## Overview

This document covers the **quarterly results** with *detailed analysis*.

### Key Metrics

| Metric | Q1 | Q2 |
|---|---|---|
| Revenue | $1M | $1.5M |
| Users | 10k | 15k |

### Action Items

1. Expand marketing
2. Hire engineers
   - Frontend
   - Backend

> Note: All figures are approximate.

```python
def calculate_growth(q1, q2):
    return (q2 - q1) / q1 * 100
```

---

## Conclusion

Results are **positive** and we should continue investing.
"""
        params = CreateDocxParams(markdown=md, filename="report.docx")
        result = await create_docx.handler(params, {})
        assert result["status"] == "created"
        assert result["size_bytes"] > 0

    @pytest.mark.asyncio
    async def test_missing_template_returns_error(self):
        params = CreateDocxParams(
            markdown="# Test",
            template_local_path="/nonexistent/template.docx",
            filename="test.docx",
        )
        result = await create_docx.handler(params, {})
        assert "error" in result
