"""Tests for the extract_docx tool."""

from __future__ import annotations

import io
from typing import Any

import pytest

from app.tools.extraction.extract_docx import ExtractDocxParams, extract_docx


def _create_minimal_docx() -> bytes:
    """Create a minimal DOCX document in memory using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph of the test document.")
    doc.add_paragraph("Second paragraph with different content.")

    # Add a simple table.
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Row 1 A"
    table.cell(1, 1).text = "Row 1 B"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocx:

    @pytest.mark.asyncio
    async def test_basic_extraction(self, tool_context: dict[str, Any]) -> None:
        """Extracting a simple DOCX should return text, tables, styles, and metadata."""
        docx_bytes = _create_minimal_docx()
        tool_context["blob"].download_blob.return_value = docx_bytes

        params = ExtractDocxParams(container="test-container", path="report.docx")
        result = await extract_docx.handler(params, tool_context)

        assert result["filename"] == "report.docx"
        assert "First paragraph" in result["text"]
        assert "Second paragraph" in result["text"]
        assert isinstance(result["tables"], list)
        assert len(result["tables"]) == 1
        assert isinstance(result["styles"], list)
        assert isinstance(result["metadata"], dict)

    @pytest.mark.asyncio
    async def test_table_structure(self, tool_context: dict[str, Any]) -> None:
        """Table extraction should produce headers and data rows."""
        docx_bytes = _create_minimal_docx()
        tool_context["blob"].download_blob.return_value = docx_bytes

        params = ExtractDocxParams(container="c", path="tbl.docx")
        result = await extract_docx.handler(params, tool_context)

        table = result["tables"][0]
        assert "headers" in table
        assert "rows" in table
        assert table["headers"] == ["Header A", "Header B"]
        assert table["rows"] == [["Row 1 A", "Row 1 B"]]

    @pytest.mark.asyncio
    async def test_blob_download_error(self, tool_context: dict[str, Any]) -> None:
        """A download failure should produce an error result, not raise."""
        tool_context["blob"].download_blob.side_effect = RuntimeError("Connection lost")

        params = ExtractDocxParams(container="c", path="bad.docx")
        result = await extract_docx.handler(params, tool_context)

        assert "error" in result
        assert "Connection lost" in result["error"]
        assert result["filename"] == "bad.docx"

    @pytest.mark.asyncio
    async def test_invalid_docx_bytes(self, tool_context: dict[str, Any]) -> None:
        """Non-DOCX bytes should return an error dict."""
        tool_context["blob"].download_blob.return_value = b"not-a-docx"

        params = ExtractDocxParams(container="c", path="corrupt.docx")
        result = await extract_docx.handler(params, tool_context)

        assert "error" in result

    @pytest.mark.asyncio
    async def test_metadata_keys(self, tool_context: dict[str, Any]) -> None:
        """Metadata should include standard DOCX properties."""
        docx_bytes = _create_minimal_docx()
        tool_context["blob"].download_blob.return_value = docx_bytes

        params = ExtractDocxParams(container="c", path="meta.docx")
        result = await extract_docx.handler(params, tool_context)

        metadata = result["metadata"]
        for key in ("author", "title", "subject", "created", "modified"):
            assert key in metadata, f"Missing metadata key: {key}"

    @pytest.mark.asyncio
    async def test_empty_docx(self, tool_context: dict[str, Any]) -> None:
        """An empty DOCX should return empty text and no tables."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        tool_context["blob"].download_blob.return_value = buf.getvalue()

        params = ExtractDocxParams(container="c", path="empty.docx")
        result = await extract_docx.handler(params, tool_context)

        assert result["text"] == ""
        assert result["tables"] == []
